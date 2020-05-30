import pandas as pd
from enum import Enum
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import json
from docx.shared import Inches


def load_products(f_name: str = 'response_1590486932629.json') -> dict:
    products = {}
    with open(f_name, encoding="utf8") as json_file:
        raw = json.load(json_file).get('products')
    for product in raw:
        products[product.get('sku')] = {'name': product.get('product_description').get('3').get('name'),
                                        'price': float(product.get('price')),
                                        }
    return products


products_t = load_products()


class OrderTypes:
    delivery = "משלוח"
    take_away = "איסוף עצמי"

    def is_delivery(self, order_type):
        if self.delivery in order_type:
            return True
        return False


class Package:
    base_r = "כלי עגול 250 גרם"
    base_b = "כלי פלסטיק 0.5 ליטר"
    u = "יחידה"

    double_r = "כלי עגול 500 גרם"
    double_b = "כלי פלסטיק 1 ליטר"

    def is_base(self, p):
        return p in [self.base_b, self.base_r, self.u]

    def package_factor(self, p: str) -> int:
        return 1 if self.is_base(p) else 2


class ColumnName(Enum):
    ORDER_NUM = "מס הזמנה"

    # header
    NAME = "שם לקוח"
    DATE = "תאריך"
    EMAIL = "אימייל"
    PHONE = "טלפון"
    TOTAL = "סה״כ תשלום"
    ORDER_TYPE = "סוג משלוח"
    ADDRESS = "תשלום כתובת"
    STATUS = "סטטוס"
    COMMENT = "הערה"
    APPROVAL_ID = "מספר אישור"
    CREDIT_NUM = "4 ספרות של הכרטיס"

    # item
    ITEM = "פריטים בהזמנה"
    PACKAGE = "אפשרויות מוצר"
    AMOUNT = "כמות פריטים"
    PRODUCT_ID = "מק''ט"


class Heb:
    status = "סטאטוס"
    delivery_method = "שיטת מסירה"
    order_num = "מספר הזמנה"
    name = "שם"
    phone = "טלפון"
    total = 'סה"כ'
    product_name = "שם מוצר"
    package = "אריזה"
    amount = "כמות"
    delivery = "משלוח"
    unit_price = 'מחיר יחידה'
    address = "כתובת"
    comment = "הערות"


class Order:
    def __init__(self, order_id: int, df: pd.DataFrame):
        self.id = order_id
        self.name = df[ColumnName.NAME.value].iloc[0]
        self.email = df[ColumnName.EMAIL.value].iloc[0]
        self.phone = df[ColumnName.PHONE.value].iloc[0]
        self.date = df[ColumnName.DATE.value].iloc[0].date()
        self.comment = df[ColumnName.COMMENT.value].iloc[0]
        self.total = df[ColumnName.TOTAL.value].iloc[0]
        self.status = df[ColumnName.STATUS.value].iloc[0]
        self.credit = df[ColumnName.CREDIT_NUM.value].iloc[0]

        self.order_type = df[ColumnName.ORDER_TYPE.value].iloc[0]
        self.address = df[ColumnName.ADDRESS.value].iloc[0]
        self.approval_id = df[ColumnName.APPROVAL_ID.value].iloc[0]

        self.produces = []
        for _, row in df.iterrows():
            self.pars_item(row)

    def pars_item(self, row):
        product_id = row[ColumnName.PRODUCT_ID.value]
        product_name = products_t.get(product_id).get('name')

        base_price = products_t.get(row[ColumnName.PRODUCT_ID.value]).get('price')
        amount = row[ColumnName.AMOUNT.value]
        package = row[ColumnName.PACKAGE.value]
        if pd.isnull(package):
            package = Package.u
        else:
            package = package.split(':')[-1]
        unit_price = base_price * Package().package_factor(package)

        total = unit_price * amount
        self.produces.append([product_name, package, unit_price, amount, total])

    def to_dict(self) -> dict:
        order = dict()
        order['title'] = f"""
                            {Heb.status}: {self.status}
                            {Heb.delivery_method}: {self.order_type}
                            {Heb.order_num}: {self.id}
                            {Heb.name}: {self.name}
                            {Heb.phone}: {self.phone}
                            """
        order["is_delivery"] = OrderTypes().is_delivery(self.order_type)
        if order["is_delivery"]:
            order['address'] = f"{Heb.address}: {self.address}"
            self.produces.append([f"{Heb.delivery}", "", "", "1", "20"])
        if not pd.isnull(self.comment):
            order['comment'] = f"{Heb.comment} : {self.comment}"
        order['total'] = self.total
        self.produces.append([f'{Heb.total}', "", "", "", f"{self.total}"])
        order['produces'] = self.produces

        return order
def set_col_widths(table):
    widths = (Inches(1), Inches(2), Inches(1.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def xl_to_doc(path: str):
    xl = pd.ExcelFile(path)
    df = xl.parse(xl.sheet_names[0], converters={ColumnName.PHONE.value: str,
                                                 })
    orders = []
    for order_id, group in df.groupby(ColumnName.ORDER_NUM.value):
        order = Order(order_id, group)
        orders.append(order)

    document = Document()

    # word doc styles:
    rtl_style = document.styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
    rtl_style.paragraph_format.alignment = 2

    font = rtl_style.font
    font.name = 'Arial'
    font.size = Pt(14)

    rtl2_style = document.styles.add_style('rtl2', WD_STYLE_TYPE.PARAGRAPH)
    rtl2_style.paragraph_format.alignment = 2
    font2 = rtl2_style.font
    font2.name = 'Arial'
    font2.size = Pt(11)
    font2.italic = True

    for order in orders:
        # creating word page for each order
        order_doc = order.to_dict()

        # title
        document.add_paragraph(order_doc.get('title'), style='rtl')

        # address
        if order_doc.get('is_delivery'):
            address = order_doc.get('address', '')
            document.add_paragraph(address, style='rtl')

        document.add_paragraph(order_doc.get('comment', ''), style='rtl')

        produces = order_doc.get('produces')
        # adding the products table
        table = document.add_table(rows=1, cols=5, style='Light Grid Accent 1')
        table.allow_autofit = False
        hdr_cells = table.rows[0].cells
        hdr_cells[4].text = Heb.product_name
        hdr_cells[3].text = Heb.package
        hdr_cells[2].text = Heb.unit_price
        hdr_cells[1].text = Heb.amount
        hdr_cells[0].text = Heb.total
        # adding style to the table header
        for i in range(5):
            hdr_cells[i].paragraphs[0].paragraph_format.alignment = 2
            hdr_cells[i].paragraphs[0].style = rtl2_style

        for product_name, package, unit_price, amount, product_sum in produces:
            row_cells = table.add_row().cells
            row_cells[0].text = str(product_sum)
            row_cells[1].text = str(amount)
            row_cells[2].text = str(unit_price)
            row_cells[3].text = package
            row_cells[4].text = product_name
            # adding style to the table row
            for i in range(5):
                row_cells[i].paragraphs[0].paragraph_format.alignment = 2
                row_cells[i].paragraphs[0].style = rtl2_style

        widths = (Inches(1), Inches(1), Inches(2), Inches(3), Inches(5))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width


        # adding page break
        document.add_page_break()
    document.save('demo.docx')


if __name__ == '__main__':
    pd.set_option('display.max_rows', 500)
    pd.set_option('display.max_columns', 500)
    pd.set_option('display.width', 1000)

    xl_to_doc("order_export_all.xlsx")
